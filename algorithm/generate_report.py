"""
맛노래 수집 데이터 통계 Word 보고서 생성
"""
import sys, io, os, sqlite3
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "crol"))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from crol_config import DB_PATH

# ── DB 데이터 수집 ──────────────────────────────────────────────
conn = sqlite3.connect(DB_PATH)
cur  = conn.cursor()

cur.execute("SELECT COUNT(*) FROM videos");        total_videos  = cur.fetchone()[0]
cur.execute("SELECT COUNT(DISTINCT channel_id) FROM channels"); total_channels = cur.fetchone()[0]
cur.execute("SELECT SUM(subscriber_count) FROM channels WHERE snapshot_at = (SELECT MAX(snapshot_at) FROM channels)")
total_subs = cur.fetchone()[0] or 0
cur.execute("SELECT MIN(snapshot_at), MAX(snapshot_at), COUNT(DISTINCT DATE(snapshot_at)) FROM videos")
mn, mx, days = cur.fetchone()
db_size = os.path.getsize(DB_PATH) / (1024*1024)

cur.execute("SELECT source, COUNT(*) FROM videos GROUP BY source")
sources = cur.fetchall()

cur.execute("SELECT keyword, COUNT(*) FROM videos WHERE keyword IS NOT NULL GROUP BY keyword ORDER BY 2 DESC LIMIT 20")
keyword_stats = cur.fetchall()

cur.execute("SELECT is_short, COUNT(*) FROM videos GROUP BY is_short")
short_stats = dict(cur.fetchall())

view_buckets = []
for label, cond in [
    ("1만 미만", "view_count < 10000"),
    ("1만~10만", "view_count BETWEEN 10000 AND 100000"),
    ("10만~100만", "view_count BETWEEN 100000 AND 1000000"),
    ("100만~1000만", "view_count BETWEEN 1000000 AND 10000000"),
    ("1000만 이상", "view_count > 10000000"),
]:
    cur.execute(f"SELECT COUNT(*) FROM videos WHERE {cond}")
    view_buckets.append((label, cur.fetchone()[0]))

ch_buckets = []
for label, cond in [
    ("1만 미만 (소형)", "subscriber_count < 10000"),
    ("1만~10만 (중형)", "subscriber_count BETWEEN 10000 AND 100000"),
    ("10만~100만 (대형)", "subscriber_count BETWEEN 100000 AND 1000000"),
    ("100만 이상 (메가)", "subscriber_count >= 1000000"),
]:
    cur.execute(f"SELECT COUNT(*) FROM channels WHERE {cond}")
    ch_buckets.append((label, cur.fetchone()[0]))

cur.execute("""
    SELECT c.title, c.subscriber_count, AVG(v.view_count) as avg_v
    FROM channels c JOIN videos v ON v.channel_id = c.channel_id
    WHERE c.snapshot_at = (SELECT MAX(snapshot_at) FROM channels c2 WHERE c2.channel_id = c.channel_id)
      AND c.subscriber_count > 100
    GROUP BY c.channel_id
    ORDER BY (avg_v / c.subscriber_count) DESC
    LIMIT 5
""")
viral_top5 = cur.fetchall()

conn.close()

# ── Word 문서 생성 ──────────────────────────────────────────────
doc = Document()

# 스타일 기본 설정
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)

# ── 제목 ───────────────────────────────────────────────────────
title = doc.add_heading("맛노래 수집 데이터 통계 보고서", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()


# ── 헬퍼 ───────────────────────────────────────────────────────
def add_table(headers: list[str], rows: list[list], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # 헤더
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
    # 바디
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for para in cells[j].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
    # 컬럼 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table


# ── 1. 수집 규모 총괄 ──────────────────────────────────────────
doc.add_heading("1. 수집 규모 총괄", level=1)
add_table(
    ["항목", "수치"],
    [
        ["수집된 영상", f"{total_videos:,}개"],
        ["추적 채널", f"{total_channels:,}개"],
        ["추적 채널 총 구독자", f"{total_subs:,}명"],
        ["수집 기간", f"{mn[:10]} ~ {mx[:10]}"],
        ["수집 일수", f"{days}일"],
        ["DB 파일 크기", f"{db_size:.2f} MB"],
    ],
    col_widths=[5, 8],
)


# ── 2. 수집 출처 ───────────────────────────────────────────────
doc.add_heading("2. 영상 수집 출처", level=1)
doc.add_paragraph(
    "YouTube Data API v3을 통해 두 가지 채널로 수집됨:\n"
    "• search: 음식 관련 키워드 + #shorts 해시태그 검색 (최근 7일)\n"
    "• popular: YouTube 한국 인기 영상 차트 (mostPopular)"
)
src_map = {"search": "키워드 검색", "popular": "인기 차트", None: "(기타/초기)"}
add_table(
    ["출처", "영상 수", "비율"],
    [
        [src_map.get(s, str(s)), f"{c:,}개", f"{c/total_videos*100:.1f}%"]
        for s, c in sources
    ],
    col_widths=[4, 4, 4],
)


# ── 3. 키워드별 수집 ───────────────────────────────────────────
doc.add_heading("3. 키워드별 수집 영상 (TOP 20)", level=1)
doc.add_paragraph("음식 카테고리·장소·트렌드 키워드 별로 수집된 쇼츠 영상 분포")
add_table(
    ["순위", "키워드", "영상 수"],
    [[i, k, f"{c:,}개"] for i, (k, c) in enumerate(keyword_stats, 1)],
    col_widths=[2, 6, 4],
)


# ── 4. 영상 형식 ───────────────────────────────────────────────
doc.add_heading("4. 영상 형식 분포", level=1)
short_n  = short_stats.get(1, 0)
normal_n = short_stats.get(0, 0)
add_table(
    ["형식", "영상 수", "비율"],
    [
        ["쇼츠 (≤60초)", f"{short_n:,}개", f"{short_n/total_videos*100:.1f}%"],
        ["일반 영상",     f"{normal_n:,}개", f"{normal_n/total_videos*100:.1f}%"],
    ],
    col_widths=[5, 4, 4],
)


# ── 5. 조회수 분포 ─────────────────────────────────────────────
doc.add_heading("5. 영상 조회수 분포", level=1)
add_table(
    ["조회수 구간", "영상 수", "비율"],
    [[lab, f"{n:,}개", f"{n/total_videos*100:.1f}%"] for lab, n in view_buckets],
    col_widths=[5, 4, 4],
)


# ── 6. 채널 구독자 분포 ────────────────────────────────────────
doc.add_heading("6. 채널 구독자 규모 분포", level=1)
total_ch = sum(n for _, n in ch_buckets)
add_table(
    ["채널 규모", "채널 수", "비율"],
    [[lab, f"{n:,}개", f"{n/max(total_ch,1)*100:.1f}%"] for lab, n in ch_buckets],
    col_widths=[6, 4, 4],
)


# ── 7. 바이럴 채널 TOP5 ────────────────────────────────────────
doc.add_heading("7. 바이럴 계수 TOP5 채널", level=1)
doc.add_paragraph(
    "바이럴 계수 = 채널 평균 영상 조회수 / 구독자 수\n"
    "구독자 수에 비해 조회수가 폭발적인 채널 — 추천 알고리즘에서 가중치를 우선 반영함"
)
add_table(
    ["순위", "채널명", "구독자", "평균 조회수", "바이럴 계수"],
    [
        [i, t[:25], f"{int(s):,}명", f"{int(a):,}뷰", f"{a/max(s,1):.1f}x"]
        for i, (t, s, a) in enumerate(viral_top5, 1)
    ],
    col_widths=[1.5, 5, 3, 3.5, 3],
)


# ── 8. 수집 메타데이터 항목 ────────────────────────────────────
doc.add_heading("8. 수집 메타데이터 항목", level=1)

doc.add_heading("8.1 videos 테이블 (영상별)", level=2)
add_table(
    ["필드", "설명"],
    [
        ["id",            "YouTube 영상 ID"],
        ["title",         "영상 제목"],
        ["description",   "영상 설명문"],
        ["tags",          "영상 태그 (JSON)"],
        ["channel",       "채널명"],
        ["channel_id",    "채널 ID (성장률 추적용)"],
        ["view_count",    "조회수"],
        ["like_count",    "좋아요 수"],
        ["comment_count", "댓글 수"],
        ["published_at",  "업로드 시각"],
        ["duration",      "영상 길이 (ISO 8601)"],
        ["is_short",      "쇼츠 여부 (1/0)"],
        ["source",        "수집 경로 (popular/search)"],
        ["keyword",       "검색 키워드"],
        ["snapshot_at",   "수집 시점"],
    ],
    col_widths=[4, 9],
)

doc.add_heading("8.2 channels 테이블 (채널별)", level=2)
add_table(
    ["필드", "설명"],
    [
        ["channel_id",       "채널 고유 ID"],
        ["title",            "채널명"],
        ["subscriber_count", "구독자 수"],
        ["video_count",      "총 영상 수"],
        ["view_count",       "총 누적 조회수"],
        ["snapshot_at",      "스냅샷 시점 (성장률 계산)"],
    ],
    col_widths=[4, 9],
)

doc.add_heading("8.3 daily_stats 테이블 (일별 집계)", level=2)
add_table(
    ["필드", "설명"],
    [
        ["date",               "집계 날짜"],
        ["top_keywords",       "당일 최다 등장 키워드"],
        ["top_title_patterns", "당일 최다 제목 2-gram"],
        ["top_tags",           "당일 최다 태그"],
        ["top_hashtags",       "당일 최다 해시태그"],
    ],
    col_widths=[5, 8],
)


# ── 9. 활용 방식 ───────────────────────────────────────────────
doc.add_heading("9. 추천 알고리즘 데이터 활용", level=1)
doc.add_paragraph(
    "수집된 데이터는 5차원 점수 합산으로 영상 검색 결과에 반영됨:"
)
add_table(
    ["스코어 항목", "가중치", "데이터 소스"],
    [
        ["키워드 매칭",    "30%", "videos.title, videos.description"],
        ["Engagement",   "30%", "(like_count + comment_count×2) / view_count"],
        ["최신성",         "20%", "videos.published_at (2년 감쇠)"],
        ["채널 성장률",    "15%", "channels.subscriber_count / 영상 평균 조회수"],
        ["조회수 정규화",  "5%",  "videos.view_count (로그)"],
    ],
    col_widths=[5, 3, 8],
)


# ── 저장 ───────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "맛노래_수집통계_보고서.docx")
doc.save(out_path)
print(f"보고서 생성 완료: {out_path}")
print(f"파일 크기: {os.path.getsize(out_path)/1024:.1f} KB")
